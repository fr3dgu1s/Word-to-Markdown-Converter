"""
Word-to-Markdown Converter — local FastAPI server.

Endpoints:
    GET    /                       Static index.html
    GET    /health                 Liveness probe
    GET    /api/status             Docling readiness
    GET    /api/update-check       Compare local checkout with GitHub main
    GET    /changelog              Static changelog
    POST   /api/convert            Single-file conversion
    POST   /api/convert-batch      Batch conversion (uploaded files)
    POST   /api/convert-folder     Scan a local folder and convert all .docx
    POST   /api/save-changes       Persist edited Markdown
    GET    /api/open-folder        Open the Outputs folder in Explorer
    POST   /api/shutdown           Stop the server
    GET    /logs/latest            Tail app.log
    DELETE /logs/latest            Truncate app.log

Command line:
    python server.py serve
    python server.py docx-to-md document.docx
    python server.py batch-docx-to-md C:\\path\\to\\folder
    python server.py md-to-docx document.md
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import time
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

from fastapi import FastAPI, Request, UploadFile, File, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse

from logging_config import setup_logging
from paths import PROJECT_ROOT, OUTPUTS_ROOT, IMAGES_ROOT, TEMP_ROOT, LOGS_ROOT

try:
    import pythoncom  # type: ignore
    import win32com.client  # type: ignore
except ImportError:
    pythoncom = None
    win32com = None

logger = setup_logging()


@asynccontextmanager
async def lifespan(_app: FastAPI):
    ensure_converter_initializing()
    threading.Thread(target=_warm_word_com, daemon=True, name="word-com-warmup").start()
    yield


app = FastAPI(title="Word-to-Markdown Converter", lifespan=lifespan)

DEFAULT_UPDATE_REPOSITORY = "fr3dgu1s/Word-to-Markdown-Converter"
DEFAULT_UPDATE_BRANCH = "main"
UPDATE_CHECK_TIMEOUT_SECONDS = 6


@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    import traceback
    logger.exception(f"Unhandled error in {request.method} {request.url.path}: {exc}")
    return JSONResponse(
        status_code=500,
        content={
            "error": str(exc),
            "traceback": traceback.format_exc(),
            "path": str(request.url),
        },
    )


# ---------------------------------------------------------------------------
# Docling lazy initialisation
# ---------------------------------------------------------------------------
_converter_ready = threading.Event()
_converter = None
_converter_error: Optional[str] = None
_word_app = None
_word_lock = threading.Lock()
_converter_init_lock = threading.Lock()
_converter_init_started = False


def get_word_app():
    """
    Return the shared persistent Word.Application COM instance.

    Word is created on first use, then kept alive for the server process so MIP
    authentication can be reused across protected-file conversions.
    """
    global _word_app
    if pythoncom is None or win32com is None:
        raise RuntimeError(
            "pywin32 is required for Purview-protected files. "
            "Install it with `python -m pip install -r requirements.txt`."
        )

    if _word_app is not None:
        try:
            _ = _word_app.Version
            return _word_app
        except Exception:
            _word_app = None

    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    _word_app = word
    logger.info("[Word COM] Word instance started and ready")
    return _word_app


def _warm_word_com() -> None:
    try:
        get_word_app()
        logger.info("[Startup] Word COM instance warmed up successfully")
    except Exception as exc:
        logger.warning(f"[Startup] Could not start Word COM: {exc}")
        logger.warning("[Startup] Non-protected files will still convert normally")
        logger.warning("[Startup] Purview-protected files will fail until Word COM is available")


def _init_converter() -> None:
    global _converter, _converter_error
    logger.info("[INIT] Starting Docling converter initialisation")
    try:
        from docling.document_converter import DocumentConverter  # noqa: PLC0415
        from docling.datamodel.pipeline_options import PdfPipelineOptions  # noqa: PLC0415
        opts = PdfPipelineOptions()
        opts.generate_picture_images = True
        opts.images_scale = 2.0
        _converter = DocumentConverter()
        logger.info("[INIT] Docling converter ready")
    except Exception as exc:
        _converter_error = str(exc)
        logger.error(f"[INIT] Docling converter failed: {exc}")
    finally:
        _converter_ready.set()


def ensure_converter_initializing() -> None:
    """Start Docling initialisation once, on demand."""
    global _converter_init_started
    if _converter_init_started:
        return
    with _converter_init_lock:
        if _converter_init_started:
            return
        _converter_init_started = True
        threading.Thread(target=_init_converter, daemon=True, name="docling-init").start()


def _get_converter():
    """Block until the converter is ready then return it (or raise on failure)."""
    ensure_converter_initializing()
    _converter_ready.wait(timeout=120)
    if _converter_error:
        raise RuntimeError(f"Document converter failed to initialise: {_converter_error}")
    if _converter is None:
        raise RuntimeError("Document converter is not available.")
    return _converter


# ---------------------------------------------------------------------------
# Naming helpers
# ---------------------------------------------------------------------------

_WINDOWS_FORBIDDEN = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def safe_md_basename(stem: str) -> str:
    """Sanitise a filename stem to be safe on Windows while preserving spaces/dots."""
    cleaned = _WINDOWS_FORBIDDEN.sub("", stem).strip().rstrip(".")
    return cleaned or "document"


def safe_image_dir(stem: str) -> str:
    """Lowercase, hyphen-separated stem used as the per-document image folder name."""
    clean = re.sub(r"[^a-zA-Z0-9]+", "-", stem).strip("-").lower()
    return clean or "document"


def batch_output_name(filename: str) -> str:
    """Return the batch output filename for an input .docx (preserves original stem)."""
    stem = Path(filename).stem
    return f"{safe_md_basename(stem)}-BATCH.md"


def single_output_name(filename: str) -> str:
    stem = Path(filename).stem
    return f"{safe_md_basename(stem)}.md"


def _unique_path(folder: Path, filename: str) -> Path:
    """If ``folder/filename`` already exists, append ``-2``, ``-3``, … to the stem."""
    candidate = folder / filename
    if not candidate.exists():
        return candidate
    stem = Path(filename).stem
    suffix = Path(filename).suffix
    counter = 2
    while True:
        candidate = folder / f"{stem}-{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def _unique_image_dir(stem_slug: str, images_root: Path = IMAGES_ROOT) -> Path:
    folder = images_root / stem_slug
    if not folder.exists():
        return folder
    counter = 2
    while True:
        candidate = images_root / f"{stem_slug}-{counter}"
        if not candidate.exists():
            return candidate
        counter += 1


# ---------------------------------------------------------------------------
# Purview / MIP protected-file helpers
# ---------------------------------------------------------------------------

def is_purview_protected(path: Path) -> bool:
    """
    A standard .docx is a valid ZIP archive.

    A Purview-encrypted .docx is a CFB (Compound File Binary) container instead,
    so Python's zipfile module raises BadZipFile.
    """
    try:
        with zipfile.ZipFile(path, "r"):
            return False
    except zipfile.BadZipFile:
        return True


def strip_protection_and_save(protected_path: Path) -> Path:
    """
    Opens a Purview-protected .docx in the persistent Word instance.
    Uses a two-step RTF intermediate to force encryption removal:
      1. Save as RTF  — RTF cannot carry MIP labels, Word strips encryption here.
      2. Re-open RTF  — now a plain, decrypted document.
      3. Save as DOCX — produces a clean .docx Docling can read.
    Avoids the SensitivityLabel COM API entirely.
    """
    rtf_temp = Path(tempfile.mktemp(suffix=".rtf"))
    docx_temp = Path(tempfile.mktemp(suffix=".docx"))

    with _word_lock:
        word = get_word_app()
        doc = None
        rtf_doc = None
        try:
            # Step 1 — Open the protected file (Word decrypts via cached MIP session)
            doc = word.Documents.Open(
                str(protected_path.resolve()),
                ReadOnly=False,
                AddToRecentFiles=False,
                Revert=False,
            )

            # Step 2 — Save as RTF (FileFormat 6 = wdFormatRTF)
            # RTF has no MIP container support; Word is forced to write plain content.
            doc.SaveAs2(
                str(rtf_temp.resolve()),
                FileFormat=6,
                AddToRecentFiles=False,
            )
            doc.Close(False)
            doc = None

            # Step 3 — Re-open the now-clean RTF
            rtf_doc = word.Documents.Open(
                str(rtf_temp.resolve()),
                ReadOnly=False,
                AddToRecentFiles=False,
            )

            # Step 4 — Save as DOCX (FileFormat 16 = wdFormatXMLDocument)
            rtf_doc.SaveAs2(
                str(docx_temp.resolve()),
                FileFormat=16,
                AddToRecentFiles=False,
            )
            rtf_doc.Close(False)
            rtf_doc = None

            return docx_temp

        except Exception as exc:
            if docx_temp.exists():
                docx_temp.unlink(missing_ok=True)
            raise RuntimeError(f"Word COM strip-via-RTF failed: {exc}") from exc

        finally:
            # Always close open documents and delete the RTF intermediate
            for d in [doc, rtf_doc]:
                if d is not None:
                    try:
                        d.Close(False)
                    except Exception:
                        pass
            if rtf_temp.exists():
                try:
                    rtf_temp.unlink()
                except Exception:
                    pass


# ---------------------------------------------------------------------------
# Update check helpers
# ---------------------------------------------------------------------------

def _run_git(args: list[str]) -> Optional[str]:
    try:
        result = subprocess.run(
            ["git", *args],
            cwd=str(PROJECT_ROOT),
            capture_output=True,
            text=True,
            timeout=5,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None
    if result.returncode != 0:
        return None
    value = result.stdout.strip()
    return value or None


def _resolve_update_repository() -> str:
    configured = os.getenv("UPDATE_CHECK_REPOSITORY", "").strip()
    if configured:
        return configured.removesuffix(".git")

    remote = _run_git(["remote", "get-url", "origin"]) or ""
    patterns = (
        r"github\.com[:/](?P<repo>[^/\s]+/[^/\s]+?)(?:\.git)?$",
        r"https?://github\.com/(?P<repo>[^/\s]+/[^/\s]+?)(?:\.git)?$",
    )
    for pattern in patterns:
        match = re.search(pattern, remote)
        if match:
            return match.group("repo")
    return DEFAULT_UPDATE_REPOSITORY


def _fetch_github_json(repository: str, path: str) -> dict:
    url = f"https://api.github.com/repos/{repository}/{path.lstrip('/')}"
    request = urllib.request.Request(
        url,
        headers={
            "Accept": "application/vnd.github+json",
            "User-Agent": "Word-to-Markdown-Converter",
        },
    )
    with urllib.request.urlopen(request, timeout=UPDATE_CHECK_TIMEOUT_SECONDS) as response:
        return json.loads(response.read().decode("utf-8"))


def _fetch_latest_github_commit(repository: str, branch: str) -> dict:
    encoded_branch = urllib.parse.quote(branch, safe="")
    return _fetch_github_json(repository, f"commits/{encoded_branch}")


def _fetch_github_compare(repository: str, local_sha: str, branch: str) -> Optional[dict]:
    encoded_local = urllib.parse.quote(local_sha, safe="")
    encoded_branch = urllib.parse.quote(branch, safe="")
    try:
        return _fetch_github_json(repository, f"compare/{encoded_local}...{encoded_branch}")
    except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, OSError) as exc:
        logger.warning(f"[UPDATE] Could not compare local commit with GitHub: {exc}")
        return None


def check_for_updates() -> dict:
    repository = _resolve_update_repository()
    branch = os.getenv("UPDATE_CHECK_BRANCH", DEFAULT_UPDATE_BRANCH).strip() or DEFAULT_UPDATE_BRANCH
    local_sha = _run_git(["rev-parse", "HEAD"])

    try:
        latest = _fetch_latest_github_commit(repository, branch)
    except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, OSError) as exc:
        logger.warning(f"[UPDATE] Could not check GitHub updates: {exc}")
        return {
            "ok": False,
            "update_available": False,
            "repository": repository,
            "branch": branch,
            "local_sha": local_sha,
            "error": str(exc),
            "checked_at": datetime.now(timezone.utc).isoformat(),
        }

    latest_sha = latest.get("sha", "")
    commit = latest.get("commit", {}) or {}
    message = (commit.get("message") or "").splitlines()[0] if commit else ""
    latest_url = latest.get("html_url") or f"https://github.com/{repository}/commit/{latest_sha}"
    compare = _fetch_github_compare(repository, local_sha, branch) if local_sha else None
    ahead_by = compare.get("ahead_by") if compare else None
    compare_status = compare.get("status") if compare else None
    update_available = (
        bool(compare and compare_status in {"ahead", "diverged"} and (ahead_by or 0) > 0)
        if compare else bool(local_sha and latest_sha and local_sha != latest_sha)
    )
    compare_url = (
        f"https://github.com/{repository}/compare/{local_sha}...{latest_sha}"
        if update_available else None
    )

    return {
        "ok": True,
        "update_available": update_available,
        "repository": repository,
        "branch": branch,
        "local_sha": local_sha,
        "local_short": local_sha[:7] if local_sha else None,
        "latest_sha": latest_sha,
        "latest_short": latest_sha[:7] if latest_sha else None,
        "latest_message": message,
        "latest_url": latest_url,
        "compare_url": compare_url,
        "ahead_by": ahead_by,
        "compare_status": compare_status,
        "checked_at": datetime.now(timezone.utc).isoformat(),
    }


# ---------------------------------------------------------------------------
# Middleware
# ---------------------------------------------------------------------------

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def log_requests(request: Request, call_next):
    start = time.time()
    logger.info(f"→ {request.method} {request.url.path}")
    response = await call_next(request)
    elapsed = round((time.time() - start) * 1000)
    logger.info(f"← {request.method} {request.url.path} {response.status_code} ({elapsed}ms)")
    return response


# ---------------------------------------------------------------------------
# Core conversion helper
# ---------------------------------------------------------------------------

def _convert_docx_path_to_markdown(
    src_path: Path,
    *,
    display_name: str,
    output_filename: str,
    include_markdown: bool = True,
    output_dir: Path = OUTPUTS_ROOT,
    images_root: Optional[Path] = None,
) -> dict:
    """Run Docling on a local ``.docx`` file and write Markdown + images.

    Writes the resulting ``.md`` under ``output_dir`` using ``output_filename``.
    Extracted images go to ``images_root/<slug>/``; by default that is
    ``output_dir/Images`` for CLI calls and ``OUTPUTS_ROOT/Images`` for the app.
    ``display_name`` is only used for logging / image-folder naming.
    """
    output_dir = Path(output_dir).expanduser().resolve()
    images_root = Path(images_root).expanduser().resolve() if images_root else output_dir / "Images"
    output_dir.mkdir(parents=True, exist_ok=True)
    images_root.mkdir(parents=True, exist_ok=True)

    original_stem = Path(display_name).stem
    image_slug = safe_image_dir(original_stem)
    image_folder = _unique_image_dir(image_slug, images_root=images_root)
    image_folder.mkdir(parents=True, exist_ok=True)

    md_file_path = _unique_path(output_dir, output_filename)
    clean_tmp: Optional[Path] = None
    docling_input_path = src_path

    t0 = time.perf_counter()
    logger.info(f"[CONVERT] start | file={display_name} -> {md_file_path.name}")

    try:
        from docling_core.types.doc import PictureItem  # noqa: PLC0415

        if is_purview_protected(src_path):
            logger.info(f"[Purview] Protected file detected: {display_name}. Stripping via Word...")
            clean_tmp = strip_protection_and_save(src_path)
            docling_input_path = clean_tmp

        conv_res = _get_converter().convert(docling_input_path)

        picture_counter = 0
        for element, _level in conv_res.document.iterate_items():
            if isinstance(element, PictureItem):
                picture_counter += 1
                img_name = f"image_{picture_counter}.png"
                with (image_folder / img_name).open("wb") as fp:
                    element.get_image(conv_res.document).save(fp, "PNG")

        raw_markdown = conv_res.document.export_to_markdown(image_placeholder="IMAGE_TOKEN")
        final_markdown = raw_markdown
        image_link_base = os.path.relpath(image_folder, md_file_path.parent).replace(os.sep, "/")
        for i in range(1, picture_counter + 1):
            tag = f"![spec-image]({image_link_base}/image_{i}.png)"
            final_markdown = final_markdown.replace("IMAGE_TOKEN", tag, 1)

        with open(md_file_path, "w", encoding="utf-8") as f:
            f.write(final_markdown)

        elapsed_ms = int((time.perf_counter() - t0) * 1000)
        logger.info(f"[CONVERT] done  | file={display_name} | elapsed={elapsed_ms}ms")

        result = {
            "doc_name": md_file_path.stem,
            "output_file": str(md_file_path),
            "image_dir": image_folder.name,
            "output_dir": str(output_dir),
        }
        if include_markdown:
            result["markdown"] = final_markdown
        return result

    except Exception as exc:
        elapsed_ms = int((time.perf_counter() - t0) * 1000)
        logger.error(f"[CONVERT] fail  | file={display_name} | elapsed={elapsed_ms}ms | error={exc}")
        raise
    finally:
        if clean_tmp and clean_tmp.exists():
            try:
                clean_tmp.unlink()
                logger.info(f"[Purview] Deleted clean temp copy: {clean_tmp}")
            except Exception as exc:
                logger.warning(f"[Purview] Could not delete clean temp copy {clean_tmp}: {exc}")


def convert_file_to_markdown(
    upload_file: UploadFile,
    *,
    output_filename: str,
    include_markdown: bool = True,
    output_dir: Path = OUTPUTS_ROOT,
) -> dict:
    """Convert one uploaded ``.docx`` file to Markdown using Docling.

    Spools the upload to a temp file, then delegates to
    :func:`_convert_docx_path_to_markdown`.
    """
    tmp_path: Optional[Path] = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=str(TEMP_ROOT)) as tmp:
            shutil.copyfileobj(upload_file.file, tmp)
            tmp_path = Path(tmp.name)

        return _convert_docx_path_to_markdown(
            tmp_path,
            display_name=upload_file.filename or "document",
            output_filename=output_filename,
            include_markdown=include_markdown,
            output_dir=output_dir,
            images_root=IMAGES_ROOT if Path(output_dir).resolve() == OUTPUTS_ROOT.resolve() else None,
        )
    finally:
        if tmp_path and tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


def find_docx_files(folder: Path, recursive: bool = True) -> list[Path]:
    """Collect .docx files in a folder, skipping Word lock files."""
    pattern = "**/*.docx" if recursive else "*.docx"
    seen: set[str] = set()
    docx_files: list[Path] = []
    for path in folder.glob(pattern):
        if not path.is_file() or path.name.startswith("~$"):
            continue
        key = str(path).lower()
        if key in seen:
            continue
        seen.add(key)
        docx_files.append(path)
    docx_files.sort()
    return docx_files


_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_BOLD_RE = re.compile(r"(\*\*|__)(.*?)\1")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)|(?<!_)_([^_\n]+)_(?!_)")
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def _plain_markdown_inline(text: str) -> str:
    """Best-effort inline Markdown cleanup for the python-docx fallback."""
    text = _IMAGE_RE.sub(lambda match: f"{match.group(1) or 'Image'} ({match.group(2)})", text)
    text = _LINK_RE.sub(lambda match: f"{match.group(1)} ({match.group(2)})", text)
    text = _BOLD_RE.sub(lambda match: match.group(2), text)
    text = _ITALIC_RE.sub(lambda match: match.group(1) or match.group(2), text)
    text = _INLINE_CODE_RE.sub(lambda match: match.group(1), text)
    return re.sub(r"\\([\\`*_{}\[\]()#+\-.!|>])", r"\1", text)


def _add_docx_paragraph(document, text: str = "", style: Optional[str] = None):
    if not style:
        return document.add_paragraph(text)
    try:
        return document.add_paragraph(text, style=style)
    except KeyError:
        return document.add_paragraph(text)


def _split_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_markdown_table_separator(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _write_markdown_table(document, table_lines: list[str]) -> None:
    rows = [_split_markdown_table_row(line) for line in table_lines]
    rows = [row for row in rows if not _is_markdown_table_separator(row)]
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            text = row[col_index] if col_index < len(row) else ""
            table.cell(row_index, col_index).text = _plain_markdown_inline(text)


def _render_markdown_with_python_docx(markdown_path: Path, output_path: Path) -> None:
    try:
        from docx import Document  # type: ignore  # noqa: PLC0415
        from docx.shared import Pt  # type: ignore  # noqa: PLC0415
    except ImportError as exc:
        raise RuntimeError(
            "Markdown-to-Word fallback requires python-docx. "
            "Install dependencies with `python -m pip install -r requirements.txt`, "
            "or install Pandoc for richer conversion."
        ) from exc

    document = Document()
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    code_lines: list[str] = []
    table_lines: list[str] = []
    in_code_block = False

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        for code_line in code_lines:
            paragraph = _add_docx_paragraph(document, style="No Spacing")
            run = paragraph.add_run(code_line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        code_lines = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            _write_markdown_table(document, table_lines)
            table_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if "|" in line and stripped.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        if not stripped:
            document.add_paragraph()
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            document.add_heading(_plain_markdown_inline(heading.group(2)), level=len(heading.group(1)))
            continue

        unordered = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if unordered:
            _add_docx_paragraph(document, _plain_markdown_inline(unordered.group(1)), style="List Bullet")
            continue

        ordered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if ordered:
            _add_docx_paragraph(document, _plain_markdown_inline(ordered.group(1)), style="List Number")
            continue

        if stripped.startswith(">"):
            quote = stripped.lstrip(">").strip()
            _add_docx_paragraph(document, _plain_markdown_inline(quote), style="Intense Quote")
            continue

        if re.fullmatch(r"[-*_]{3,}", stripped):
            document.add_paragraph("-----")
            continue

        document.add_paragraph(_plain_markdown_inline(line))

    flush_table()
    flush_code()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _resolve_docx_output_path(markdown_path: Path, output_path: Optional[Path]) -> Path:
    if output_path is None:
        return _unique_path(OUTPUTS_ROOT, f"{safe_md_basename(markdown_path.stem)}.docx")

    output_path = output_path.expanduser()
    if output_path.exists() and output_path.is_dir():
        return output_path / f"{safe_md_basename(markdown_path.stem)}.docx"
    if output_path.suffix.lower() != ".docx":
        return output_path / f"{safe_md_basename(markdown_path.stem)}.docx"
    return output_path


def convert_markdown_path_to_docx(
    markdown_path: Path,
    *,
    output_path: Optional[Path] = None,
    use_pandoc: bool = True,
) -> dict:
    markdown_path = markdown_path.expanduser().resolve()
    if not markdown_path.exists() or not markdown_path.is_file():
        raise FileNotFoundError(f"Markdown file does not exist: {markdown_path}")
    if markdown_path.suffix.lower() not in {".md", ".markdown", ".txt"}:
        raise ValueError("Markdown-to-Word input must be a .md, .markdown, or .txt file.")

    target = _resolve_docx_output_path(markdown_path, output_path).resolve()
    target.parent.mkdir(parents=True, exist_ok=True)

    pandoc_path = shutil.which("pandoc") if use_pandoc else None
    if pandoc_path:
        result = subprocess.run(
            [pandoc_path, str(markdown_path), "-o", str(target)],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError((result.stderr or result.stdout or "Pandoc conversion failed.").strip())
        method = "pandoc"
    else:
        _render_markdown_with_python_docx(markdown_path, target)
        method = "python-docx"

    return {
        "input": str(markdown_path),
        "output_file": str(target),
        "method": method,
    }


# ---------------------------------------------------------------------------
# Health / status
# ---------------------------------------------------------------------------

@app.get("/health")
async def health_check():
    return {"ok": True}


@app.get("/api/status")
async def converter_status():
    ensure_converter_initializing()
    ready = _converter_ready.is_set() and _converter is not None
    return {"converter_ready": ready, "error": _converter_error}


@app.get("/api/update-check")
async def update_check():
    return check_for_updates()


# ---------------------------------------------------------------------------
# Static / utility
# ---------------------------------------------------------------------------

@app.get("/")
async def serve_index():
    return FileResponse("index.html")


@app.get("/changelog")
async def serve_changelog():
    changelog_path = PROJECT_ROOT / "CHANGELOG.md"
    if not changelog_path.exists():
        raise HTTPException(status_code=404, detail="CHANGELOG.md was not found.")
    return FileResponse(changelog_path, media_type="text/markdown")


app.mount("/Outputs", StaticFiles(directory=str(OUTPUTS_ROOT)), name="Outputs")


@app.get("/api/open-folder")
async def open_folder():
    os.startfile(OUTPUTS_ROOT)
    return {"status": "opened", "folder": str(OUTPUTS_ROOT)}


@app.post("/api/shutdown")
async def shutdown_app():
    logger.info("[SHUTDOWN] Shutdown requested by local user")

    def stop_server():
        time.sleep(1)
        os._exit(0)

    threading.Thread(target=stop_server, daemon=True).start()
    return {"status": "shutting_down"}


@app.post("/api/save-changes")
async def save_changes(data: dict = Body(...)):
    output_file = data.get("output_file")
    doc_name = data.get("doc_name")
    content = data.get("markdown", "")

    if output_file:
        target = Path(output_file)
        # Pin the write inside OUTPUTS_ROOT to prevent path traversal.
        try:
            target.resolve().relative_to(OUTPUTS_ROOT.resolve())
        except ValueError:
            raise HTTPException(status_code=400, detail="output_file must be inside Outputs.")
    elif doc_name:
        target = OUTPUTS_ROOT / f"{safe_md_basename(doc_name)}.md"
    else:
        raise HTTPException(status_code=400, detail="output_file or doc_name is required.")

    target.parent.mkdir(parents=True, exist_ok=True)
    with open(target, "w", encoding="utf-8") as f:
        f.write(content)
    return {"status": "saved", "path": str(target)}


# ---------------------------------------------------------------------------
# Local conversion endpoints
# ---------------------------------------------------------------------------

@app.post("/api/convert")
async def convert_document(file: UploadFile = File(...)):
    filename = file.filename or ""
    if not filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    try:
        result = convert_file_to_markdown(
            file,
            output_filename=single_output_name(filename),
            include_markdown=True,
        )
        return {
            "markdown": result["markdown"],
            "doc_name": result["doc_name"],
            "output_file": result["output_file"],
            "output_dir": str(OUTPUTS_ROOT),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Unhandled error in /api/convert: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/convert-batch")
async def convert_documents_batch(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="No files were provided.")

    converted_files: list[dict] = []
    failed_files: list[dict] = []

    for upload_file in files:
        filename = upload_file.filename or ""
        if not filename.lower().endswith(".docx"):
            failed_files.append({
                "input": filename,
                "error": "Unsupported file type (only .docx is accepted)",
            })
            continue
        try:
            item = convert_file_to_markdown(
                upload_file,
                output_filename=batch_output_name(filename),
                include_markdown=False,
            )
            converted_files.append({
                "input": filename,
                "output": item["output_file"],
            })
        except Exception as exc:
            logger.exception(f"Unhandled error in /api/convert-batch for {filename}: {exc}")
            failed_files.append({
                "input": filename,
                "error": str(exc) or "Document could not be read",
            })

    return {
        "output_dir": str(OUTPUTS_ROOT),
        "converted_count": len(converted_files),
        "failed_count": len(failed_files),
        "converted_files": converted_files,
        "failed_files": failed_files,
    }


@app.post("/api/convert-folder")
async def convert_documents_in_folder(data: dict = Body(...)):
    """Scan a local folder for ``.docx`` files and convert each one.

    Body: ``{"folder_path": "C:/path/to/folder", "recursive": true}``
    The server reads files directly from disk (no upload) — only safe in
    this app because it binds to ``127.0.0.1``.
    """
    raw_path = (data.get("folder_path") or "").strip().strip('"').strip("'")
    if not raw_path:
        raise HTTPException(status_code=400, detail="folder_path is required.")
    recursive = bool(data.get("recursive", True))

    try:
        folder = Path(raw_path).expanduser().resolve()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Invalid folder path: {exc}")

    if not folder.exists() or not folder.is_dir():
        raise HTTPException(status_code=400, detail=f"Folder does not exist: {folder}")

    docx_files = find_docx_files(folder, recursive=recursive)
    logger.info(f"[FOLDER-SCAN] folder={folder} recursive={recursive} found={len(docx_files)}")

    if not docx_files:
        return {
            "output_dir": str(OUTPUTS_ROOT),
            "scanned_folder": str(folder),
            "scanned_count": 0,
            "converted_count": 0,
            "failed_count": 0,
            "converted_files": [],
            "failed_files": [],
        }

    converted_files: list[dict] = []
    failed_files: list[dict] = []

    for src in docx_files:
        try:
            item = _convert_docx_path_to_markdown(
                src,
                display_name=src.name,
                output_filename=batch_output_name(src.name),
                include_markdown=False,
            )
            converted_files.append({"input": str(src), "output": item["output_file"]})
        except Exception as exc:
            logger.exception(f"Unhandled error in /api/convert-folder for {src}: {exc}")
            failed_files.append({
                "input": str(src),
                "error": str(exc) or "Document could not be read",
            })

    return {
        "output_dir": str(OUTPUTS_ROOT),
        "scanned_folder": str(folder),
        "scanned_count": len(docx_files),
        "converted_count": len(converted_files),
        "failed_count": len(failed_files),
        "converted_files": converted_files,
        "failed_files": failed_files,
    }


# ---------------------------------------------------------------------------
# Log viewer endpoints
# ---------------------------------------------------------------------------

@app.get("/logs/latest")
def logs_latest(lines: int = 100):
    lines = min(max(lines, 1), 500)
    log_path = LOGS_ROOT / "app.log"
    if not log_path.exists():
        return {"lines": []}
    with open(log_path, encoding="utf-8", errors="replace") as f:
        all_lines = f.readlines()
    return {"lines": [ln.rstrip("\n") for ln in all_lines[-lines:]]}


@app.delete("/logs/latest")
def logs_clear():
    log_path = LOGS_ROOT / "app.log"
    if log_path.exists():
        open(log_path, "w").close()
    logger.info("[LOGS] Log file cleared by user")
    return {"status": "cleared"}


def _print_result(result: dict, *, as_json: bool) -> None:
    if as_json:
        print(json.dumps(result, indent=2))
        return

    if "converted_files" in result:
        print(
            f"Converted {result['converted_count']} of {result['scanned_count']} document(s) "
            f"to {result['output_dir']}"
        )
        for item in result["converted_files"]:
            print(f"OK: {item['input']} -> {item['output']}")
        for item in result["failed_files"]:
            print(f"FAILED: {item['input']} | {item['error']}", file=sys.stderr)
        return

    if "method" in result:
        print(f"Converted {result['input']} -> {result['output_file']} ({result['method']})")
        return

    print(f"Converted {result.get('input', result.get('doc_name', 'document'))} -> {result['output_file']}")


def _run_cli_docx_to_md(args: argparse.Namespace) -> int:
    source = Path(args.input).expanduser().resolve()
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(f"Word document does not exist: {source}")
    if source.suffix.lower() != ".docx":
        raise ValueError("DOCX-to-Markdown input must be a .docx file.")

    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else OUTPUTS_ROOT
    output_name = args.output_name or single_output_name(source.name)
    result = _convert_docx_path_to_markdown(
        source,
        display_name=source.name,
        output_filename=output_name,
        include_markdown=args.print_markdown,
        output_dir=output_dir,
        images_root=IMAGES_ROOT if output_dir == OUTPUTS_ROOT.resolve() else None,
    )
    result["input"] = str(source)

    if args.print_markdown:
        print(result["markdown"])
    else:
        _print_result(result, as_json=args.json)
    return 0


def _run_cli_batch_docx_to_md(args: argparse.Namespace) -> int:
    folder = Path(args.folder).expanduser().resolve()
    if not folder.exists() or not folder.is_dir():
        raise NotADirectoryError(f"Folder does not exist: {folder}")

    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else OUTPUTS_ROOT
    docx_files = find_docx_files(folder, recursive=args.recursive)
    converted_files: list[dict] = []
    failed_files: list[dict] = []

    for source in docx_files:
        try:
            item = _convert_docx_path_to_markdown(
                source,
                display_name=source.name,
                output_filename=batch_output_name(source.name),
                include_markdown=False,
                output_dir=output_dir,
                images_root=IMAGES_ROOT if output_dir == OUTPUTS_ROOT.resolve() else None,
            )
            converted_files.append({"input": str(source), "output": item["output_file"]})
        except Exception as exc:
            logger.exception(f"CLI batch conversion failed for {source}: {exc}")
            failed_files.append({"input": str(source), "error": str(exc) or "Document could not be read"})

    result = {
        "output_dir": str(output_dir),
        "scanned_folder": str(folder),
        "scanned_count": len(docx_files),
        "converted_count": len(converted_files),
        "failed_count": len(failed_files),
        "converted_files": converted_files,
        "failed_files": failed_files,
    }
    _print_result(result, as_json=args.json)
    return 1 if failed_files else 0


def _run_cli_md_to_docx(args: argparse.Namespace) -> int:
    result = convert_markdown_path_to_docx(
        Path(args.input),
        output_path=Path(args.output).expanduser() if args.output else None,
        use_pandoc=not args.no_pandoc,
    )
    _print_result(result, as_json=args.json)
    return 0


def _run_cli_convert(args: argparse.Namespace) -> int:
    source = Path(args.input).expanduser().resolve()
    if source.is_dir():
        args.folder = str(source)
        args.recursive = args.recursive
        return _run_cli_batch_docx_to_md(args)
    if source.suffix.lower() == ".docx":
        args.output_name = None
        args.print_markdown = False
        return _run_cli_docx_to_md(args)
    if source.suffix.lower() in {".md", ".markdown", ".txt"}:
        args.output = args.output_dir
        args.no_pandoc = False
        return _run_cli_md_to_docx(args)
    raise ValueError("Input must be a .docx, .md/.markdown/.txt file, or a folder containing .docx files.")


def _run_server(host: str = "127.0.0.1", port: int = 8000) -> int:
    import uvicorn  # noqa: PLC0415

    uvicorn.run(app, host=host, port=port)
    return 0


def build_cli_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Run the Markdown Studio server or convert Word/Markdown files from the command line."
    )
    subparsers = parser.add_subparsers(dest="command")

    serve = subparsers.add_parser("serve", help="Start the local FastAPI web server.")
    serve.add_argument("--host", default="127.0.0.1", help="Host interface to bind. Default: 127.0.0.1")
    serve.add_argument("--port", type=int, default=8000, help="TCP port to bind. Default: 8000")
    serve.set_defaults(func=lambda args: _run_server(args.host, args.port))

    docx_to_md = subparsers.add_parser(
        "docx-to-md",
        aliases=["word-to-md"],
        help="Convert one .docx file to Markdown.",
    )
    docx_to_md.add_argument("input", help="Path to the .docx file.")
    docx_to_md.add_argument("-o", "--output-dir", help="Directory for the generated .md and Images folder.")
    docx_to_md.add_argument("--output-name", help="Generated Markdown filename. Default: input stem + .md")
    docx_to_md.add_argument("--print-markdown", action="store_true", help="Print converted Markdown to stdout.")
    docx_to_md.add_argument("--json", action="store_true", help="Print machine-readable JSON.")
    docx_to_md.set_defaults(func=_run_cli_docx_to_md)

    batch = subparsers.add_parser(
        "batch-docx-to-md",
        aliases=["batch"],
        help="Convert every .docx in a folder to Markdown.",
    )
    batch.add_argument("folder", help="Folder to scan for .docx files.")
    batch.add_argument("-o", "--output-dir", help="Directory for generated .md files and Images folder.")
    batch.add_argument("--recursive", dest="recursive", action="store_true", default=True, help="Scan subfolders.")
    batch.add_argument("--no-recursive", dest="recursive", action="store_false", help="Only scan the top folder.")
    batch.add_argument("--json", action="store_true", help="Print machine-readable JSON.")
    batch.set_defaults(func=_run_cli_batch_docx_to_md)

    md_to_docx = subparsers.add_parser(
        "md-to-docx",
        aliases=["md-to-word"],
        help="Convert one Markdown file to a .docx file.",
    )
    md_to_docx.add_argument("input", help="Path to the .md, .markdown, or .txt file.")
    md_to_docx.add_argument("-o", "--output", help="Output .docx path or destination folder.")
    md_to_docx.add_argument("--no-pandoc", action="store_true", help="Use the built-in python-docx fallback.")
    md_to_docx.add_argument("--json", action="store_true", help="Print machine-readable JSON.")
    md_to_docx.set_defaults(func=_run_cli_md_to_docx)

    convert = subparsers.add_parser(
        "convert",
        help="Auto-detect input type: .docx to Markdown, Markdown to .docx, or folder batch conversion.",
    )
    convert.add_argument("input", help="Input .docx, Markdown file, or folder.")
    convert.add_argument("-o", "--output-dir", help="Output directory, or output .docx path for Markdown input.")
    convert.add_argument("--recursive", dest="recursive", action="store_true", default=True, help="Scan subfolders.")
    convert.add_argument("--no-recursive", dest="recursive", action="store_false", help="Only scan the top folder.")
    convert.add_argument("--json", action="store_true", help="Print machine-readable JSON.")
    convert.set_defaults(func=_run_cli_convert)

    return parser


def main(argv: Optional[list[str]] = None) -> int:
    argv = list(sys.argv[1:] if argv is None else argv)
    if not argv:
        return _run_server()

    parser = build_cli_parser()
    args = parser.parse_args(argv)
    if not hasattr(args, "func"):
        parser.print_help()
        return 2

    try:
        if getattr(args, "json", False):
            for handler in logging.getLogger().handlers:
                if isinstance(handler, logging.StreamHandler) and not isinstance(handler, logging.FileHandler):
                    handler.setLevel(logging.WARNING)
        return args.func(args)
    except KeyboardInterrupt:
        return 130
    except Exception as exc:
        logger.exception(f"Command failed: {exc}")
        print(f"Error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
